from pathlib import Path


def extract_text(file_path: str) -> str:
    path = Path(file_path)
    ext = path.suffix.lower()
    try:
        if ext == ".docx":
            return _extract_docx(path)
        elif ext in (".xlsx", ".xls"):
            return _extract_excel(path)
        elif ext == ".pdf":
            return _extract_pdf(path)
    except Exception:
        return ""
    return ""


def _extract_docx(path: Path) -> str:
    from docx import Document

    doc = Document(str(path))
    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    return "\n".join(parts)


def _extract_excel(path: Path) -> str:
    ext = path.suffix.lower()
    parts: list[str] = []
    if ext == ".xlsx":
        from openpyxl import load_workbook

        wb = load_workbook(str(path), read_only=True, data_only=True)
        for ws in wb.worksheets:
            for row in ws.iter_rows():
                for cell in row:
                    if cell.value is not None:
                        text = str(cell.value).strip()
                        if text:
                            parts.append(text)
        wb.close()
    else:
        import xlrd

        wb = xlrd.open_workbook(str(path))
        for ws in wb.sheets():
            for row_idx in range(ws.nrows):
                for col_idx in range(ws.ncols):
                    val = ws.cell_value(row_idx, col_idx)
                    text = str(val).strip()
                    if text:
                        parts.append(text)
    return "\n".join(parts)


def _extract_pdf(path: Path) -> str:
    import pdfplumber

    parts: list[str] = []
    with pdfplumber.open(str(path)) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                parts.append(text)
    return "\n".join(parts)
